import io
import re

from docx import Document
from docx.shared import Inches

try:
    from image_scrapper import downloader
except ImportError:
    from .image_scrapper import downloader


async def generate_docx_prompt(language, emotion_type, topic):
    message = f"""Create an {language} language very long outline for a {emotion_type} research paper on the topic of {topic} which is long as much as possible. 
Language of research paper - {language}.
Provide as much information as possible.

Put this tag before the Title: [TITLE]
Put this tag after the Title: [/TITLE]
Put this tag before the Subtitle: [SUBTITLE]
Put this tag after the Subtitle: [/SUBTITLE]
Put this tag before the Heading: [HEADING]
Put this tag after the Heading: [/HEADING]
Put this tag before the Content: [CONTENT]
Put this tag after the Content: [/CONTENT]
Put this tag before the Image: [IMAGE]
Put this tag after the Image: [/IMAGE]

Elaborate on the Content, provide as much information as possible.
You put a [/CONTENT] at the end of the Content.
Do not put a tag before ending previous.

For example:
[TITLE]Mental Health[/TITLE]
[SUBTITLE]Understanding and Nurturing Your Mind: A Guide to Mental Health[/SUBTITLE]
[HEADING]Mental Health Definition[/HEADING]
[CONTENT]...[/CONTENT]
[IMAGE]Person Meditating[/IMAGE]

Pay attention to the language of research paper - {language}.
Each image should be described in general by a set of keywords, such as "Mount Everest Sunset" or "Niagara Falls Rainbow".
Do not reply as if you are talking about the research paper itself. (ex. "Include pictures here about...")
Do not include any special characters (?, !, ., :, ) in the Title.
Do not include any additional information in your response and stick to the format."""

    return message


async def generate_docx(answer):
    doc = Document()

    async def split_tags(reply):
        pattern = r'\[(.*?)\](.*?)\[/\1\]'
        tags = re.findall(pattern, reply, re.DOTALL)
        return tags

    async def parse_response(tags_array):
        if not tags_array:
            raise IndexError
        for item in tags_array:
            match (item[0]):
                case('TITLE'):
                    doc.add_heading(item[1], 0)
                case('SUBTITLE'):
                    doc.add_heading(item[1], 1)
                case('HEADING'):
                    doc.add_heading(item[1], 2)
                case('CONTENT'):
                    doc.add_paragraph(item[1])
                case('IMAGE'):
                    try:
                        image_data = await downloader.download(item[1], limit=1, adult_filter_off=True, timeout=15,
                                                               filter="+filterui:aspect-wide+filterui:imagesize-wallpaper+filterui:photo-photo")
                        doc.add_picture(io.BytesIO(image_data), width=Inches(6))
                    except Exception:
                        pass

    async def find_title(tags_array):
        for item in tags_array:
            if item[0] == 'TITLE':
                return item[1]

    reply_array = await split_tags(answer)
    await parse_response(reply_array)
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    docx_title = f"{await find_title(reply_array)}.docx"
    print(f"done {docx_title}")

    return docx_bytes, docx_title
